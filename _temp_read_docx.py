import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

out_lines = []

doc = Document(r'C:\Users\prpau\OneDrive\Desktop\RepositoriosGit\Contrato-automatico\CPS EMPRESTIMO BLINDSAFE.docx')

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        out_lines.append(f'[P{i:03d}] {t[:150]}')

out_lines.append('')

for ti, table in enumerate(doc.tables):
    out_lines.append(f'--- Tabela {ti} ({len(table.rows)}x{len(table.columns)}) ---')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            t = cell.text.strip()
            if t:
                out_lines.append(f'  [{ri},{ci}] {t[:120]}')

output = '\n'.join(out_lines)

with open(r'C:\Users\prpau\OneDrive\Desktop\RepositoriosGit\Contrato-automatico\_temp_docx_output.txt', 'w', encoding='utf-8') as f:
    f.write(output)

print("Done. Lines written:", len(out_lines))
