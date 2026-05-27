"""
make_veiculo_template.py
Gera templates/contrato_veiculo.docx a partir do documento original.
Substitui os campos em vermelho por marcadores Jinja2.
Execute sempre que o texto do documento for atualizado.
"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC  = Path(r"C:\Users\prpau\OneDrive\Desktop\Contrato-rápido\contrato-blindsafe\CONTRATO COMPRA DE VEÍCULO.docx")
DEST = Path(__file__).parent / "templates" / "contrato_veiculo.docx"

# Mapeamento: texto vermelho → placeholder Jinja2
# Runs com textos que se combinam (ex: 'Nº' + ' DO RG') são tratados abaixo
RED_MAP = {
    'NOME COMPLETO DO CONTRATANTE': '{{ nome }}',
    'nacionalidade':                '{{ nacionalidade }}',
    '000.000.000-00':               '{{ cpf }}',
    'ENDEREÇO COMPLETO, BAIRRO, CIDADE/UF, CEP': '{{ endereco_completo }}',
    'EMAIL@DOMINIO.COM':            '{{ email }}',
    'MARCA/MODELO DO VEÍCULO':      '{{ veiculo_modelo }}',
    'MARCA/MODELO DO VEÍCULO':      '{{ veiculo_modelo }}',
    'AAAA/AAAA':                    '{{ veiculo_ano }}',
    'COR':                          '{{ veiculo_cor }}',
    'NÚMERO DO CHASSI':             '{{ veiculo_chassi }}',
    'PLACA':                        '{{ veiculo_placa }}',
    'RENAVAM':                      '{{ veiculo_renavam }}',
}

DATE_PARA_IDX = 39  # "Rio de Janeiro, DD de MÊS de AAAA."


def _is_red(run) -> bool:
    color_el = run._element.find('.//' + qn('w:color'))
    if color_el is None:
        return False
    val = color_el.get(qn('w:val'), '')
    return val.upper() == 'FF0000'


def _set_color_auto(run):
    """Remove a cor explícita do run para ficar preta como o texto normal."""
    color_el = run._element.find('.//' + qn('w:color'))
    if color_el is not None:
        color_el.set(qn('w:val'), '000000')


def _process_paragraph(para):
    """Substitui grupos de runs vermelhos pelos placeholders correspondentes."""
    runs = list(para.runs)
    i = 0
    while i < len(runs):
        r = runs[i]
        if not _is_red(r):
            i += 1
            continue

        # Coleta runs vermelhos consecutivos
        group = [r]
        j = i + 1
        while j < len(runs) and _is_red(runs[j]):
            group.append(runs[j])
            j += 1

        combined = ''.join(rr.text for rr in group).strip()

        # Tenta encontrar o placeholder pelo texto combinado
        placeholder = RED_MAP.get(combined)

        # Fallback: tenta com o texto do primeiro run sozinho
        if placeholder is None:
            placeholder = RED_MAP.get(group[0].text.strip())

        # Caso especial: "Nº" + " DO RG" → rg
        if placeholder is None and 'RG' in combined:
            placeholder = '{{ rg }}'

        # Caso especial: "MARCA/MODELO " + "DO VEÍCULO" (split em 8.2)
        if placeholder is None and 'MODELO' in combined and 'VEÍCULO' in combined:
            placeholder = '{{ veiculo_modelo }}'

        if placeholder is not None:
            group[0].text = placeholder
            _set_color_auto(group[0])
            for rr in group[1:]:
                rr.text = ''
                _set_color_auto(rr)
        else:
            # Se não mapeado, só remove a cor vermelha
            for rr in group:
                _set_color_auto(rr)

        i = j


def _replace_date_paragraph(para):
    """Substitui o parágrafo de data inteiro por {{ local_data }}."""
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    run = para.add_run('{{ local_data }}')
    run.bold = False


def main():
    if not SRC.exists():
        print(f"Arquivo original não encontrado: {SRC}")
        return

    doc = Document(str(SRC))

    for i, para in enumerate(doc.paragraphs):
        if i == DATE_PARA_IDX:
            _replace_date_paragraph(para)
        else:
            _process_paragraph(para)

    DEST.parent.mkdir(exist_ok=True)
    doc.save(str(DEST))
    print(f"Template salvo em: {DEST}")

    # Verificação
    doc2 = Document(str(DEST))
    print("\nPlaceholders encontrados no template:")
    for i, p in enumerate(doc2.paragraphs):
        if '{{' in p.text:
            print(f"  [{i:03d}] {p.text[:100]}")


if __name__ == "__main__":
    main()
