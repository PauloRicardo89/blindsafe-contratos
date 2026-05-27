"""
make_procuracao_template.py
Gera templates/procuracao.docx (PF) e templates/procuracao_pj.docx (PJ)
a partir do documento fonte com as duas procurações (Dr. Leo + BlindSafe).
Execute sempre que o texto fixo do documento for atualizado.
"""
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = Path(r"C:\Users\prpau\OneDrive\Desktop\Contrato-rápido\contrato-blindsafe\Procuração e Hipo\PROCURAÇÃO DR LEO E BLINDSAFE.docx")
DEST_PF = Path(__file__).parent / "templates" / "procuracao.docx"
DEST_PJ = Path(__file__).parent / "templates" / "procuracao_pj.docx"

# Índices dos parágrafos a substituir (baseado na estrutura do documento fonte)
OUTORGANTE_IDXS = [2, 29]   # parágrafo com o bloco do outorgante
DATE_IDXS       = [8, 35]   # parágrafo com a data
NOME_IDXS       = [11, 38]  # parágrafo com o nome na assinatura


def _clear_runs(para):
    for r in list(para.runs):
        r._element.getparent().remove(r._element)


def _add_run(para, text, bold=None):
    from docx.shared import Pt
    run = para.add_run(text)
    if bold is not None:
        run.bold = bold
    run.font.name = "Tahoma"
    run.font.size = Pt(11)
    return run


def _set_outorgante_pf(para):
    _clear_runs(para)
    _add_run(para, "OUTORGANTE", bold=True)
    _add_run(para, ": ", bold=True)
    _add_run(para, "{{r bloco_procuracao }}")


def _set_outorgante_pj(para):
    _clear_runs(para)
    _add_run(para, "OUTORGANTE", bold=True)
    _add_run(para, ": ", bold=True)
    _add_run(para, "{{r bloco_procuracao_pj }}  ")
    _add_run(para, "OUTORGANTE", bold=True)
    _add_run(para, ": ", bold=True)
    _add_run(para, "{{r bloco_procuracao }}")


def _set_local_data(para):
    _clear_runs(para)
    _add_run(para, "{{ local_data }}")


def _set_nome(para):
    _clear_runs(para)
    _add_run(para, "{{ nome }}", bold=True)


def _remove_inline_breaks(para):
    """Remove <w:br> elements embutidos dentro dos runs do parágrafo."""
    for br in para._element.findall('.//' + qn('w:br')):
        br.getparent().remove(br)


def _add_page_break_before(para):
    """Adiciona quebra de página antes do parágrafo via pPr/pageBreakBefore."""
    pPr = para._element.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)


def _remove_blank_paragraphs(doc, start_idx, end_idx):
    """Remove parágrafos em branco no intervalo [start_idx, end_idx)."""
    body = doc.element.body
    # Coleta os elementos a remover (de trás pra frente para não deslocar índices)
    paras = doc.paragraphs
    to_remove = [paras[i]._element for i in range(start_idx, end_idx)
                 if not paras[i].text.strip()]
    for el in to_remove:
        body.remove(el)


def build_template(outorgante_fn, dest: Path):
    doc = Document(str(SRC))

    # Remove os parágrafos em branco (12–26) entre as duas procurações
    # e limpa os \n\n\n\n embutidos no título da segunda procuração
    _remove_blank_paragraphs(doc, 12, 27)

    # Após remoção os índices dos parágrafos restantes mudaram — recalcula
    paras = doc.paragraphs
    # Encontra os dois títulos "PROCURAÇÃO" e os dois OUTORGANTE, datas e nomes
    titulo_idxs      = [i for i, p in enumerate(paras) if 'PROCURA' in p.text and len(p.text.strip()) < 20]
    outorgante_idxs  = [i for i, p in enumerate(paras) if p.text.strip().startswith('OUTORGANTE')]
    date_idxs        = [i for i, p in enumerate(paras) if p.text.strip().startswith('Rio de Janeiro') or '{{ local_data }}' in p.text]
    nome_idxs        = [i for i, p in enumerate(paras) if i > 5 and p.text.strip() and not p.text.strip().startswith('_') and not any(k in p.text for k in ('OUTORG', 'PODER', 'Rio', 'PROCURA', '{{'))]

    # Adiciona quebra de página antes do segundo título PROCURAÇÃO
    if len(titulo_idxs) >= 2:
        segundo_titulo = paras[titulo_idxs[1]]
        _remove_inline_breaks(segundo_titulo)
        _add_page_break_before(segundo_titulo)

    for i in outorgante_idxs:
        outorgante_fn(paras[i])

    for i in date_idxs:
        _set_local_data(paras[i])

    # Nome (assinatura): parágrafos após a linha de traços "__..."
    traco_idxs = [i for i, p in enumerate(paras) if p.text.strip().startswith('___')]
    for i in traco_idxs:
        if i + 1 < len(paras):
            _set_nome(paras[i + 1])

    dest.parent.mkdir(exist_ok=True)
    doc.save(str(dest))
    print(f"Salvo: {dest}")


def main():
    if not SRC.exists():
        print(f"Arquivo fonte não encontrado: {SRC}")
        return

    build_template(_set_outorgante_pf, DEST_PF)
    build_template(_set_outorgante_pj, DEST_PJ)
    print("\nTemplates gerados com sucesso.")


if __name__ == "__main__":
    main()
